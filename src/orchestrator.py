"""
orchestrator.py — Phase 3 update

Run-aware, formatting-preserving write-back.
Supports both preset path (truth_path=None → uses preset_report.truth.json)
and upload path (truth_path=<path to generated truth.json>).
Includes diagram_inserter hook (Part D.5).
"""

import os
import time
from typing import TypedDict, Callable, Optional
from langgraph.graph import StateGraph, START, END
from groq import Groq
from docx import Document
from docx.shared import Pt
import shutil
from docx.enum.text import WD_ALIGN_PARAGRAPH
from logger import log
from truth_loader import load_truth

# --- Formatting helpers ------------------------------------------------------

_ALIGN_ENUM = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def apply_formatting(run, paragraph, format_dict: dict):
    """Apply every field from a formatting dict to a run + paragraph.
    Never leave anything to a default — fully explicit."""
    if format_dict.get("bold") is not None:
        run.bold = bool(format_dict["bold"])
    if format_dict.get("italic") is not None:
        run.italic = bool(format_dict["italic"])
    if format_dict.get("font"):
        run.font.name = format_dict["font"]
    if format_dict.get("size_pt"):
        run.font.size = Pt(format_dict["size_pt"])
    if format_dict.get("underline") is not None:
        run.underline = bool(format_dict["underline"])

    alignment = format_dict.get("alignment", "left")
    paragraph.alignment = _ALIGN_ENUM.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)


def _clear_paragraph_text(p):
    """Surgically clear only <w:t> (text) nodes from a paragraph's runs,
    leaving <w:br> (page/line breaks) and all other XML elements intact.
    This preserves manual page breaks that are embedded in runs.
    We blank the text instead of removing the <w:t> tag to prevent docx-preview crashes on empty runs."""
    from docx.oxml.ns import qn
    for t_elem in p._element.xpath('.//w:t'):
        t_elem.text = ''
    
    # Strip numbering to prevent empty bullets showing up
    for numPr in p._element.xpath('.//w:pPr/w:numPr'):
        numPr.getparent().remove(numPr)
        
    # If it's explicitly styled as a list, remove the style
    try:
        if p.style and 'List' in p.style.name:
            p.style = p.part.document.styles['Normal']
    except Exception:
        pass


def _strip_images(doc: Document):
    """Remove all floating/inline image XML elements from the document so the
    web viewer doesn't render them. Strips <w:drawing>, <w:pict>, and empty <mc:AlternateContent>."""
    from docx.oxml.ns import qn
    
    # Custom namespace for markup compatibility
    mc_ns = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    
    tags_to_strip = [
        qn('w:drawing'),
        qn('w:pict'),
        f'{{{mc_ns}}}AlternateContent',
        '{urn:schemas-microsoft-com:vml}shape'
    ]
    
    for tag in tags_to_strip:
        for elem in doc.element.body.iter(tag):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

def get_all_paragraphs(doc: Document) -> list:
    """Walk the document in XML order (same as template_parser) to get a flat list
    of all paragraphs and table-cell paragraphs, so we can index into it safely."""
    paragraphs = []
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph
            paragraphs.append(Paragraph(child, doc))
        elif tag == "tbl":
            from docx.table import Table
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        paragraphs.append(para)
    return paragraphs


def write_paragraph_node(doc: Document, node: dict, text: str = "", existing_p=None) -> tuple:
    """
    Add a paragraph for a node, applying all formatting from the node's
    formatting block (and per-run formatting if runs are present).

    If existing_p is provided, it clears that paragraph and replaces its runs
    in-place instead of creating a new paragraph.
    Returns (paragraph, run) for the body run (or the first run for multi-span).
    """
    if existing_p is not None:
        p = existing_p
        _clear_paragraph_text(p)  # Safe clear — preserves page breaks
    else:
        p = doc.add_paragraph()
    fmt = node.get("formatting", {})
    node_runs = node.get("runs", [])

    if node_runs and text == "":
        # Multi-span node: write each span with its own formatting
        for span in node_runs:
            run = p.add_run(span.get("text", ""))
            span_fmt = {
                "bold": span.get("bold"),
                "italic": span.get("italic"),
                "font": span.get("font"),
                "size_pt": span.get("size_pt"),
                "underline": span.get("underline"),
                "alignment": fmt.get("alignment", "left"),
            }
            apply_formatting(run, p, span_fmt)
        return p, p.runs[0] if p.runs else p.add_run("")
    else:
        # Single-span node (heading or body during streaming)
        run = p.add_run(text)
        apply_formatting(run, p, fmt)
        return p, run


# --- Safe save ---------------------------------------------------------------

def _safe_save(doc: Document, path: str):
    """Atomic write via temp file with retries for Windows file locks."""
    temp_path = path + ".tmp"
    for attempt in range(10):
        try:
            doc.save(temp_path)
            os.replace(temp_path, path)
            break
        except PermissionError:
            if attempt == 9:
                log.error("Safe save failed: Access denied after 10 retries.")
            time.sleep(0.2)
        except Exception as e:
            log.error(f"Safe save failed: {e}")
            break
    if os.path.exists(temp_path):
        try:
            os.remove(temp_path)
        except Exception:
            pass


# --- LangGraph state ---------------------------------------------------------

class AgentState(TypedDict):
    topic: str
    docx_path: str
    current_section_index: int
    generated_text: str
    stream_chat_callback: Callable[[str], None]
    section_callback: Callable[[str], None]
    doc_updated_callback: Callable[[], None]
    stats_callback: Callable[[dict], None]
    truth_nodes: list
    section_groups: list
    is_template: bool
    original_para_count: int
    mode: str
    target_sections: list[str]
    words_generated: int
    diagrams_generated: int
    start_time: float


# --- Graph node --------------------------------------------------------------

def generate_and_write_section(state: AgentState):
    idx = state["current_section_index"]
    section_name = state["section_groups"][idx]

    # Get nodes for this section (heading + body only; skip figure/metadata)
    nodes = [n for n in state["truth_nodes"] if n.get("section") == section_name]
    heading_node = next((n for n in nodes if n["role"] == "section_heading"), None)
    body_node = next((n for n in nodes if n["role"] == "body_paragraph"), None)

    topic = state["topic"]
    chat_callback = state["stream_chat_callback"]
    section_callback = state.get("section_callback")
    docx_path = state["docx_path"]

    log.info(f"Generating section {idx + 1}/{len(state['section_groups'])}: {section_name}")
    
    if section_callback:
        section_callback(section_name)

    doc_text = ""

    try:
        if not os.environ.get("GROQ_API_KEY"):
            raise ValueError("GROQ_API_KEY environment variable not set.")

        # Load current document
        doc = Document(docx_path) if os.path.exists(docx_path) else Document()
        
        # Build the index map if we are using an uploaded template
        all_paras = get_all_paragraphs(doc) if state.get("is_template") else []
        shift = len(all_paras) - state.get("original_para_count", 0) if state.get("is_template") else 0

        # Write heading using run-aware helper
        if heading_node:
            p_idx = heading_node.get("position", {}).get("index")
            if p_idx is not None:
                p_idx += shift
            existing = all_paras[p_idx] if all_paras and p_idx is not None and p_idx < len(all_paras) else None
            write_paragraph_node(doc, heading_node, text=section_name, existing_p=existing)

        # Prepare first body paragraph (empty, will be filled by streaming)
        body_existing = None
        if body_node:
            p_idx = body_node.get("position", {}).get("index")
            if p_idx is not None:
                p_idx += shift
            body_existing = all_paras[p_idx] if all_paras and p_idx is not None and p_idx < len(all_paras) else None
            
        body_p, body_run = write_paragraph_node(doc, body_node or {}, text="", existing_p=body_existing)
        
        # Clear any subsequent body paragraphs in this section so template lorem ipsum is removed
        all_body_nodes = [n for n in nodes if n["role"] == "body_paragraph"]
        if len(all_body_nodes) > 1:
            for extra_node in all_body_nodes[1:]:
                p_idx = extra_node.get("position", {}).get("index")
                if p_idx is not None:
                    p_idx += shift
                if all_paras and p_idx is not None and p_idx < len(all_paras):
                    _clear_paragraph_text(all_paras[p_idx])  # Safe clear — preserves page breaks
                    
        # Generate text unless in diagrams mode
        if state["mode"] != "diagrams":
            client = Groq()
            prompt = (
                f"You are writing a section for a document. "
                f"The section is '{section_name}' and the topic is '{topic}'. "
                f"First, write a single short sentence explaining what you are writing "
                f"(this will go to the chat). Then, type exactly '---' on a new line. "
                f"Finally, write the actual concise academic paragraph "
                f"(this will go into the document)."
            )

            stream = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.1-8b-instant",
                stream=True,
            )

            full_buffer = ""
            hit_delimiter = False
            chunk_count = 0
            
            for chunk in stream:
                text = chunk.choices[0].delta.content
                if text:
                    full_buffer += text

                    if not hit_delimiter:
                        if "---" in full_buffer:
                            hit_delimiter = True
                            parts = full_buffer.split("---", 1)
                            doc_text = parts[1].lstrip("\n")
                            if chat_callback:
                                chat_callback("\n")
                        else:
                            if chat_callback:
                                chat_callback(text)
                    else:
                        doc_text += text
                        if body_run:
                            body_run.text = doc_text
                        chunk_count += 1
                        
                        if state.get("stats_callback"):
                            state["stats_callback"]({
                                "words": state["words_generated"] + len(doc_text.split()),
                                "diagrams": state["diagrams_generated"],
                                "elapsed_seconds": time.time() - state["start_time"]
                            })

                        if chunk_count % 4 == 0 and state.get("doc_updated_callback"):
                            _safe_save(doc, docx_path)
                            state["doc_updated_callback"]()

            state["words_generated"] += len(doc_text.split())
            # Save after text generation
            _safe_save(doc, docx_path)
            if state.get("doc_updated_callback"):
                state["doc_updated_callback"]()
        else:
            # In diagram mode, we don't generate new text. We use the existing text for context.
            doc_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # --- Phase 5: Diagram Generation (Optional per section) ---
        # We generate diagrams in full_generation and diagrams mode.
        # In edit mode, we generally don't generate diagrams unless specifically requested, but it's safe to try.
        try:
            from diagram_agent import generate_and_insert_diagram
            # Note: diagram_inserter saves the document internally if it created it, but we passed ours so we need to save it.
            inserted = generate_and_insert_diagram(section_name, doc_text, docx_path, doc=doc, target_paragraph=body_p)
            if inserted:
                _safe_save(doc, docx_path)
                state["diagrams_generated"] += 1
                if state.get("stats_callback"):
                    state["stats_callback"]({
                        "words": state["words_generated"],
                        "diagrams": state["diagrams_generated"],
                        "elapsed_seconds": time.time() - state["start_time"]
                    })
                log.info(f"Diagram inserted for section '{section_name}'")
        except Exception as diagram_err:
            log.warning(f"diagram_inserter failed for '{section_name}': {diagram_err}")

    except Exception as e:
        log.error(f"Error generating section {section_name}: {e}")
        doc_text = f"[Content generation failed for {section_name}: {e}]"

    return {"current_section_index": idx + 1, "generated_text": doc_text}


def router(state: AgentState):
    if state["current_section_index"] < len(state["section_groups"]):
        return "generate_and_write_section"
    return END


# Build graph
workflow = StateGraph(AgentState)
workflow.add_node("generate_and_write_section", generate_and_write_section)
workflow.add_conditional_edges(START, router)  # Route at START to handle empty section_groups gracefully
workflow.add_conditional_edges("generate_and_write_section", router)
app = workflow.compile()


# --- Entry point -------------------------------------------------------------

def run_full_generation(
    topic: str,
    docx_path: str,
    stream_chat_callback: Callable[[str], None] = None,
    section_callback: Callable[[str], None] = None,
    doc_updated_callback: Callable[[], None] = None,
    stats_callback: Callable[[dict], None] = None,
    truth_path: Optional[str] = None,
):
    """
    Run full document generation.

    truth_path: path to a truth.json file.
      - None → uses preset templates/preset_report.truth.json
      - str  → uses the upload-path generated truth.json at that path
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))

    if truth_path is None:
        truth_path = os.path.join(
            os.path.dirname(current_dir), "templates", "preset_report.truth.json"
        )

    truth_nodes, source_template = load_truth(truth_path)

    # Collect unique generatable sections (skip figure/metadata/boilerplate and TOC entries)
    section_groups = []
    for node in truth_nodes:
        section = node.get("section")
        role = node.get("role", "")
        # Filter out TOC entries
        if section and "\t" not in section and section.lower().strip() not in ("contents", "table of contents") and role in ("section_heading", "body_paragraph") and section not in section_groups:
            section_groups.append(section)

    if not section_groups:
        raise ValueError(
            "This template could not be parsed into sections. "
            "The upload pipeline may have classified all paragraphs as 'body_paragraph' with no section headings. "
            "Please try re-uploading the template."
        )

    # Classify intent
    from intent_classifier import classify_intent
    intent = classify_intent(topic, section_groups)
    mode = intent["mode"]
    target_sections = intent["target_sections"]

    if stream_chat_callback:
        if mode == "edit":
            stream_chat_callback(f"\n[Agent]: Edit mode detected. Targeting sections: {', '.join(target_sections) if target_sections else 'None'}\n\n")
        elif mode == "diagrams":
            stream_chat_callback("\n[Agent]: Diagram mode detected. Generating diagrams only.\n\n")
        else:
            stream_chat_callback("\n[Agent]: Full generation mode detected.\n\n")

    # If edit mode, filter section groups to only those targeted
    if mode == "edit" and target_sections:
        # Fuzzy match target sections to actual sections
        matched_sections = []
        for target in target_sections:
            for sec in section_groups:
                if target.lower() in sec.lower() or sec.lower() in target.lower():
                    if sec not in matched_sections:
                        matched_sections.append(sec)
        if matched_sections:
            section_groups = matched_sections
        else:
            log.warning(f"Could not match target sections {target_sections} to any available sections. Defaulting to full generation.")
            mode = "full_generation"

    is_template = False
    original_para_count = 0

    if mode == "full_generation":
        # Reset output docx for a clean run
        if os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except Exception:
                pass

        if source_template and os.path.exists(source_template):
            # We have an uploaded template! Copy it so we preserve page size, margins, and page borders.
            shutil.copy2(source_template, docx_path)
            is_template = True
            # Strip all images from the copy so the web viewer renders cleanly.
            _img_doc = Document(docx_path)
            _strip_images(_img_doc)
            _safe_save(_img_doc, docx_path)
            original_para_count = len(get_all_paragraphs(_img_doc))
        else:
            # Fallback to default blank document
            _safe_save(Document(), docx_path)
    else:
        # Edit or diagram mode: use existing docx_path as-is
        is_template = True
        if os.path.exists(docx_path):
            _img_doc = Document(docx_path)
            original_para_count = len(get_all_paragraphs(_img_doc))
        else:
            # If for some reason working.docx doesn't exist, we must fallback
            _safe_save(Document(), docx_path)

    initial_state = AgentState(
        topic=topic,
        docx_path=docx_path,
        current_section_index=0,
        generated_text="",
        stream_chat_callback=stream_chat_callback,
        section_callback=section_callback,
        doc_updated_callback=doc_updated_callback,
        stats_callback=stats_callback,
        truth_nodes=truth_nodes,
        section_groups=section_groups,
        is_template=is_template,
        original_para_count=original_para_count,
        mode=mode,
        target_sections=target_sections,
        words_generated=0,
        diagrams_generated=0,
        start_time=time.time()
    )

    log.info(f"Starting orchestration — topic: {topic!r}, mode: {mode}, sections: {section_groups}")
    final_state = app.invoke(initial_state)
    log.info("Orchestration complete.")
    return final_state
