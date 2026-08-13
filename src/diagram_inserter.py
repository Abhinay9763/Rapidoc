import os
from docx import Document
from docx.shared import Inches

def insert_diagram(docx_path: str, image_path: str, caption: str, after_section: str, doc: Document = None, target_paragraph = None):
    """
    Inserts a diagram into the docx document immediately after the matching section heading,
    followed by a caption.
    """
    try:
        should_save = False
        if doc is None:
            doc = Document(docx_path)
            should_save = True
        
        if target_paragraph is None:
            for p in doc.paragraphs:
                if after_section.lower() in p.text.lower():
                    target_paragraph = p
                    break
                    
        if not target_paragraph:
            print(f"Section heading '{after_section}' not found in {docx_path}")
            return
            
        # To insert *after*, we can insert before the *next* paragraph
        next_p = target_paragraph._element.getnext()
        if next_p is not None:
            new_p = doc.add_paragraph()
            # Move the new paragraph before the next one
            next_p.addprevious(new_p._element)
            run = new_p.add_run()
            run.add_picture(image_path, width=Inches(5.5))
            
            caption_p = doc.add_paragraph()
            caption_p.style = 'Caption' if 'Caption' in doc.styles else 'Normal'
            run_cap = caption_p.add_run(caption)
            run_cap.italic = True
            next_p.addprevious(caption_p._element)
        else:
            # End of document
            doc.add_picture(image_path, width=Inches(5.5))
            p_cap = doc.add_paragraph()
            run_cap = p_cap.add_run(caption)
            run_cap.italic = True

        if should_save:
            doc.save(docx_path)
    except Exception as e:
        print(f"Failed to insert diagram: {e}")
