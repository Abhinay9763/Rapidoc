"""
template_parser.py — Phase 3, Part A

Parses an uploaded .docx template into a flat list of paragraph/table-cell
records, with Word's incidental run-splitting normalized away before any
further processing.
"""

import os
import re
import shutil
import zipfile
import tempfile
from typing import Optional
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from logger import log


_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    None: "left",
}


def _sanitize_docx(src_path: str) -> str:
    """
    WPS Office sometimes writes Target="NULL" in relationship files, which
    causes python-docx to crash with 'There is no item named NULL in the archive'.

    This function copies the docx to a temp file, strips NULL targets from all
    .rels XML files inside the zip, and returns the path to the clean copy.
    The caller is responsible for deleting the temp file.
    """
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()

    _NULL_PATTERN = re.compile(
        r'<Relationship\b[^>]*Target="NULL"[^>]*/>', re.IGNORECASE
    )

    with zipfile.ZipFile(src_path, "r") as zin, \
         zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".rels"):
                text = data.decode("utf-8", errors="replace")
                cleaned = _NULL_PATTERN.sub("", text)
                if cleaned != text:
                    log.warning(f"Stripped NULL relationship(s) from {item.filename}")
                data = cleaned.encode("utf-8")
            zout.writestr(item, data)

    return tmp.name


def _pt_from_size(size) -> Optional[float]:
    """Convert python-docx Pt/Emu size object to plain float pt, or None."""
    if size is None:
        return None
    try:
        return round(size.pt, 1)
    except AttributeError:
        return None


def _run_fingerprint(run) -> tuple:
    """Return a hashable formatting fingerprint for a single run."""
    font = run.font
    return (
        bool(run.bold or font.bold),
        bool(run.italic or font.italic),
        font.name or "",
        _pt_from_size(font.size),
        bool(font.underline),
    )


def _merge_runs(paragraph) -> list[dict]:
    """
    Merge adjacent runs with identical formatting fingerprints.
    Returns a list of merged-run dicts:
      { "text": str, "bold": bool, "italic": bool,
        "font": str, "size_pt": float|None, "underline": bool }
    """
    merged: list[dict] = []

    for run in paragraph.runs:
        if not run.text:
            continue

        fp = _run_fingerprint(run)
        span = {
            "text": run.text,
            "bold": fp[0],
            "italic": fp[1],
            "font": fp[2],
            "size_pt": fp[3],
            "underline": fp[4],
            "_fp": fp,
        }

        if merged and merged[-1]["_fp"] == fp:
            # Same formatting as last span — just extend it
            merged[-1]["text"] += run.text
        else:
            merged.append(span)

    # Strip the internal _fp key before returning
    for span in merged:
        span.pop("_fp", None)

    return merged


def _alignment_str(paragraph) -> str:
    return _ALIGN_MAP.get(paragraph.alignment, "left")


def _paragraph_to_record(para, index: int, source: str = "body") -> dict:
    """Convert a python-docx paragraph to a normalized record dict."""
    merged = _merge_runs(para)

    if not merged:
        # Empty paragraph — still a valid node (acts as spacing)
        return {
            "type": "paragraph",
            "source": source,
            "position": {"index": index},
            "style_name": para.style.name if para.style else "Normal",
            "alignment": _alignment_str(para),
            "text": "",
            "formatting": {
                "bold": False,
                "italic": False,
                "font": "",
                "size_pt": None,
                "underline": False,
                "alignment": _alignment_str(para),
            },
            "runs": [],
        }

    alignment = _alignment_str(para)
    style_name = para.style.name if para.style else "Normal"

    if len(merged) == 1:
        # Single homogeneous span — flatten into one formatting block
        span = merged[0]
        return {
            "type": "paragraph",
            "source": source,
            "position": {"index": index},
            "style_name": style_name,
            "alignment": alignment,
            "text": span["text"],
            "formatting": {
                "bold": span["bold"],
                "italic": span["italic"],
                "font": span["font"],
                "size_pt": span["size_pt"],
                "underline": span["underline"],
                "alignment": alignment,
            },
            "runs": [],
        }
    else:
        # Multi-span: dominant span = longest one
        dominant = max(merged, key=lambda s: len(s["text"]))
        full_text = "".join(s["text"] for s in merged)
        return {
            "type": "paragraph",
            "source": source,
            "position": {"index": index},
            "style_name": style_name,
            "alignment": alignment,
            "text": full_text,
            "formatting": {
                "bold": dominant["bold"],
                "italic": dominant["italic"],
                "font": dominant["font"],
                "size_pt": dominant["size_pt"],
                "underline": dominant["underline"],
                "alignment": alignment,
            },
            "runs": merged,
        }


def parse_template(docx_path: str) -> list[dict]:
    """
    Parse an uploaded .docx template into a flat list of normalized records.

    Each record has:
      type, source, position, style_name, alignment, text, formatting, runs

    Table cells are included as type="table_cell" with a source="table".
    Returns records in document order.
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Template file not found: {docx_path}")

    log.info(f"Parsing template: {docx_path}")

    # Sanitize WPS NULL-relationship issues before handing to python-docx
    clean_path = _sanitize_docx(docx_path)
    try:
        doc = Document(clean_path)
    finally:
        try:
            os.remove(clean_path)
        except Exception:
            pass

    records: list[dict] = []
    global_index = 0

    # Walk document body in XML order so paragraphs and tables are interleaved
    # correctly — python-docx's doc.paragraphs and doc.tables are separate lists
    # that lose position relative to each other.
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Regular paragraph
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            rec = _paragraph_to_record(para, global_index, source="body")
            records.append(rec)
            global_index += 1

        elif tag == "tbl":
            # Table — walk cells
            from docx.table import Table
            table = Table(child, doc)
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for cell_para in cell.paragraphs:
                        rec = _paragraph_to_record(cell_para, global_index, source="table")
                        rec["type"] = "table_cell"
                        rec["table_position"] = {
                            "row": row_idx,
                            "col": col_idx,
                        }
                        records.append(rec)
                        global_index += 1

    log.info(f"Parsed {len(records)} records from template "
             f"({sum(1 for r in records if r['type'] == 'paragraph')} paragraphs, "
             f"{sum(1 for r in records if r['type'] == 'table_cell')} table cells)")
    return records
